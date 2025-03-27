"""Document processing utilities for term sheet validation"""

import os
import logging
from werkzeug.utils import secure_filename
import PyPDF2
import docx
import pandas as pd
import re
from PIL import Image

from utils.ocr_utils import perform_ocr
from utils.nlp_utils import extract_entities, extract_dates, extract_amounts

logger = logging.getLogger(__name__)

def process_document(file_path, file_type):
    """
    Process a document and extract relevant information based on file type
    
    Args:
        file_path (str): Path to the document file
        file_type (str): Type of the document (pdf, docx, xlsx, etc.)
    
    Returns:
        dict: Extracted data from the document
    """
    try:
        logger.debug(f"Processing document: {file_path} of type {file_type}")
        
        # Extract text based on file type
        extracted_text = extract_text_from_document(file_path, file_type)
        
        if not extracted_text:
            logger.error(f"Failed to extract text from {file_path}")
            return None
        
        # Process the extracted text using NLP
        extracted_data = process_text(extracted_text)
        
        return extracted_data
    
    except Exception as e:
        logger.error(f"Error processing document {file_path}: {str(e)}")
        return None

def extract_text_from_document(file_path, file_type):
    """
    Extract text from a document based on its file type
    
    Args:
        file_path (str): Path to the document file
        file_type (str): Type of the document (pdf, docx, xlsx, etc.)
    
    Returns:
        str: Extracted text from the document
    """
    file_type = file_type.lower()
    
    try:
        # PDF processing
        if file_type == 'pdf':
            return extract_text_from_pdf(file_path)
        
        # Word document processing
        elif file_type in ['docx', 'doc']:
            return extract_text_from_word(file_path)
        
        # Excel processing
        elif file_type in ['xlsx', 'xls']:
            return extract_text_from_excel(file_path)
        
        # Image processing with OCR
        elif file_type in ['png', 'jpg', 'jpeg']:
            return extract_text_from_image(file_path)
        
        else:
            logger.error(f"Unsupported file type: {file_type}")
            return None
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return None

def extract_text_from_pdf(file_path):
    """Extract text from PDF files"""
    text = ""
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Check if PDF is empty
            if len(pdf_reader.pages) == 0:
                return ""
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        # If no text was extracted, the PDF might be scanned
        if not text.strip():
            logger.debug("PDF appears to be scanned. Attempting OCR.")
            return perform_ocr(file_path)
        
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        return None

def extract_text_from_word(file_path):
    """Extract text from Word documents"""
    text = ""
    
    try:
        doc = docx.Document(file_path)
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from Word document {file_path}: {str(e)}")
        return None

def extract_text_from_excel(file_path):
    """Extract text from Excel spreadsheets"""
    text = ""
    
    try:
        # Read the Excel file
        excel_file = pd.ExcelFile(file_path)
        
        # Process each sheet
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name)
            
            # Convert column names to string
            text += " ".join(str(col) for col in df.columns) + "\n"
            
            # Process each row
            for _, row in df.iterrows():
                # Convert row values to string and join
                text += " ".join(str(val) for val in row.values) + "\n"
        
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from Excel file {file_path}: {str(e)}")
        return None

def extract_text_from_image(file_path):
    """Extract text from images using OCR"""
    try:
        # Perform OCR on the image
        return perform_ocr(file_path)
    
    except Exception as e:
        logger.error(f"Error extracting text from image {file_path}: {str(e)}")
        return None

def process_text(text):
    """
    Process extracted text to identify key information using NLP
    
    Args:
        text (str): Text extracted from a document
    
    Returns:
        dict: Structured data extracted from the text
    """
    if not text:
        return {}
    
    extracted_data = {}
    
    # Extract entities (organizations, people, locations)
    entities = extract_entities(text)
    if entities.get('ORG'):
        extracted_data['parties'] = entities.get('ORG')
    
    # Extract dates
    dates = extract_dates(text)
    if dates:
        # Assign dates based on context - assuming first date is effective date
        if len(dates) >= 1:
            extracted_data['effective_date'] = dates[0]
        if len(dates) >= 2:
            extracted_data['expiration_date'] = dates[1]
    
    # Extract monetary amounts
    amounts = extract_amounts(text)
    if amounts:
        extracted_data['transaction_amount'] = amounts[0] if amounts else None
    
    # Extract interest rate using regex
    interest_pattern = r'(?:interest\s+rate|rate\s+of\s+interest)(?:\s+of)?(?:\s+is)?\s*\:?\s*(\d+(?:\.\d+)?)\s*%'
    interest_matches = re.findall(interest_pattern, text.lower())
    if interest_matches:
        extracted_data['interest_rate'] = interest_matches[0]
    
    # Extract payment terms - this is complex and may need refinement
    payment_pattern = r'(?:payment\s+terms|payment\s+schedule)(?:\s*\:|\s+are|\s+is)\s*([^\.]+)'
    payment_matches = re.findall(payment_pattern, text.lower())
    if payment_matches:
        extracted_data['payment_terms'] = payment_matches[0].strip()
    
    # Extract governing law
    law_pattern = r'(?:governing\s+law|jurisdiction)(?:\s*\:|\s+is|\s+shall\s+be)\s*([^\.;]+)'
    law_matches = re.findall(law_pattern, text.lower())
    if law_matches:
        extracted_data['governing_law'] = law_matches[0].strip()
    
    # Extract collateral if present
    collateral_pattern = r'(?:collateral|security)(?:\s*\:|\s+is|\s+shall\s+be)\s*([^\.;]+)'
    collateral_matches = re.findall(collateral_pattern, text.lower())
    if collateral_matches:
        extracted_data['collateral'] = collateral_matches[0].strip()
    
    return extracted_data
